from docx import Document
import json
import re

# Read the Word document
doc = Document(r'C:\Users\ZENDYGAME\Desktop\KISI-KISI TO TKA MATEMATIKA SMP KAB. BOGOR TIPE B.docx')

# Extract all text
all_text = []
for para in doc.paragraphs:
    if para.text.strip():
        all_text.append(para.text.strip())

# Also check tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                all_text.append(cell.text.strip())

# Save to JSON for analysis
with open('extracted_content.json', 'w', encoding='utf-8') as f:
    json.dump(all_text, f, ensure_ascii=False, indent=2)

print(f"Extracted {len(all_text)} lines")
print("\nFirst 30 lines:")
for i, text in enumerate(all_text[:30], 1):
    print(f"{i}. {text}")
