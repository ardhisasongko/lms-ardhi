from docx import Document
import json
import re

# Read the Word document
doc = Document(r'C:\Users\ZENDYGAME\Desktop\KISI-KISI TO TKA MATEMATIKA SMP KAB. BOGOR TIPE B.docx')

# Extract all text
all_text = []
for para in doc.paragraphs:
    if para.text.strip():
        all_text.append(para.text.strip())

# Also check tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                all_text.append(cell.text.strip())

# Parse the kisi-kisi into structured data
questions_data = []
current_question = {}

# Skip header lines (first 11 lines are headers)
content = all_text[11:]

i = 0
while i < len(content):
    line = content[i].strip()
    
    # Check if this is a question number (1-30)
    if line.isdigit() and 1 <= int(line) <= 30:
        # Save previous question if exists
        if current_question:
            questions_data.append(current_question)
        
        # Start new question
        current_question = {
            'number': int(line),
            'answer_key': '',
            'element': '',
            'sub_element': '',
            'cognitive_level': '',
            'context': '',
            'competency': '',
            'indicator': ''
        }
        
        # Next line should be answer key
        if i + 1 < len(content):
            current_question['answer_key'] = content[i + 1].strip()
            i += 2
            
            # Collect the next few lines as metadata
            metadata_lines = []
            while i < len(content) and not (content[i].strip().isdigit() and 1 <= int(content[i].strip()) <= 30):
                metadata_lines.append(content[i].strip())
                i += 1
                if len(metadata_lines) >= 6:  # We expect 6 metadata fields
                    break
            
            # Assign metadata
            if len(metadata_lines) >= 6:
                current_question['element'] = metadata_lines[0]
                current_question['sub_element'] = metadata_lines[1] if len(metadata_lines) > 1 else ''
                current_question['cognitive_level'] = metadata_lines[2] if len(metadata_lines) > 2 else ''
                current_question['context'] = metadata_lines[3] if len(metadata_lines) > 3 else ''
                current_question['competency'] = metadata_lines[4] if len(metadata_lines) > 4 else ''
                current_question['indicator'] = metadata_lines[5] if len(metadata_lines) > 5 else ''
            
            continue
    
    i += 1

# Add last question
if current_question:
    questions_data.append(current_question)

# Save parsed data
with open('parsed_kisi_kisi.json', 'w', encoding='utf-8') as f:
    json.dump(questions_data, f, ensure_ascii=False, indent=2)

print(f"Parsed {len(questions_data)} questions")
print("\nSample question:")
if questions_data:
    sample = questions_data[0]
    print(json.dumps(sample, ensure_ascii=False, indent=2))
